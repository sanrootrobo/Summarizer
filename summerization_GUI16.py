import os
import time
import subprocess
import google.generativeai as genai
import tkinter as tk
from tkinter import ttk, scrolledtext, filedialog, messagebox
import threading
import queue
import logging
from concurrent.futures import ThreadPoolExecutor, as_completed
from typing import Optional, Dict, List, Tuple, Set
import re
import argparse
from time import sleep
import json
from datetime import datetime
import html2text
import random
import shutil
import webbrowser
import tempfile
import sys
import collections

# --- Backend Standard Library Imports ---
import requests
from pathlib import Path
from urllib.parse import urljoin, urlparse, unquote
import yaml

# --- Dependency Availability Flags ---
try:
    from playwright.sync_api import sync_playwright, TimeoutError as PlaywrightTimeoutError
    PLAYWRIGHT_AVAILABLE = True
except ImportError:
    PLAYWRIGHT_AVAILABLE = False
try:
    import yt_dlp
    YT_DLP_AVAILABLE = True
except ImportError:
    YT_DLP_AVAILABLE = False

# --- Backend Third-Party Imports ---
try:
    from bs4 import BeautifulSoup
    from langchain_google_genai import ChatGoogleGenerativeAI
    from langchain_core.prompts import ChatPromptTemplate
    from langchain_core.output_parsers import StrOutputParser
    import fitz  # PyMuPDF
    import docx
    DEPENDENCIES_AVAILABLE = True
except ImportError as e:
    DEPENDENCIES_AVAILABLE = False
    print(f"Warning: Some critical dependencies are missing: {e}. The application will run in demo mode.")

# --- Markdown Preview Dependencies ---
try:
    import markdown
    from markdown.extensions import codehilite, tables, fenced_code, toc
    MARKDOWN_AVAILABLE = True
except ImportError:
    MARKDOWN_AVAILABLE = False
    print("Warning: markdown package not available. Install with: pip install markdown")

# --- Dark Theme Configuration ---
DARK_THEME = {
    'bg': '#1e1e1e',
    'fg': '#ffffff',
    'select_bg': '#264f78',
    'select_fg': '#ffffff',
    'insert_bg': '#ffffff',
    'frame_bg': '#2d2d2d',
    'button_bg': '#404040',
    'button_fg': '#ffffff',
    'entry_bg': '#3c3c3c',
    'entry_fg': '#ffffff',
    'text_bg': '#1e1e1e',
    'text_fg': '#d4d4d4',
    'scrollbar_bg': '#404040',
    'scrollbar_fg': '#686868',
    'accent': '#0078d4',
    'success': '#16c60c',
    'warning': '#ffcc02',
    'error': '#e74856',
    'border': '#404040'
}

# --- USER_AGENTS constant for the Playwright researcher ---
USER_AGENTS = [
    "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36",
    "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36",
    "Mozilla/5.0 (Windows NT 10.0; Win64; x64; rv:121.0) Gecko/20100101 Firefox/121.0",
    "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/605.1.15 (KHTML, like Gecko) Version/17.2 Safari/605.1.15",
]

class DarkThemeManager:
    """Manages dark theme styling for the application"""

    @staticmethod
    def configure_dark_theme(root):
        """Configure dark theme for the entire application, building on native themes."""
        style = ttk.Style(root)
        try:
            if sys.platform == "win32": style.theme_use('vista')
            elif sys.platform == "darwin": style.theme_use('aqua')
            else: style.theme_use('clam')
        except tk.TclError: style.theme_use('clam')
        style.configure('TFrame', background=DARK_THEME['frame_bg'])
        style.configure('TLabel', background=DARK_THEME['frame_bg'], foreground=DARK_THEME['fg'])
        style.configure('TLabelFrame', background=DARK_THEME['frame_bg'], foreground=DARK_THEME['fg'])
        style.configure('TLabelFrame.Label', background=DARK_THEME['frame_bg'], foreground=DARK_THEME['fg'])
        style.configure('TButton', background=DARK_THEME['button_bg'], foreground=DARK_THEME['button_fg'], borderwidth=1, focuscolor='none')
        style.map('TButton', background=[('active', DARK_THEME['accent']), ('pressed', DARK_THEME['select_bg'])])
        style.configure('Accent.TButton', background=DARK_THEME['accent'], foreground='white', font=('Segoe UI', 10, 'bold'))
        style.map('Accent.TButton', background=[('active', '#106ebe'), ('pressed', '#005a9e')])
        style.configure('TEntry', fieldbackground=DARK_THEME['entry_bg'], foreground=DARK_THEME['entry_fg'], insertcolor=DARK_THEME['insert_bg'])
        style.configure('TNotebook', background=DARK_THEME['frame_bg'], borderwidth=0)
        style.configure('TNotebook.Tab', background=DARK_THEME['button_bg'], foreground=DARK_THEME['fg'], padding=[12, 8], font=('Segoe UI', 9))
        style.map('TNotebook.Tab', background=[('selected', DARK_THEME['accent']), ('active', DARK_THEME['select_bg'])])
        style.configure('TCheckbutton', background=DARK_THEME['frame_bg'], foreground=DARK_THEME['fg'], focuscolor='none')
        style.configure('TRadiobutton', background=DARK_THEME['frame_bg'], foreground=DARK_THEME['fg'], focuscolor='none')
        style.configure('TScale', background=DARK_THEME['frame_bg'], troughcolor=DARK_THEME['button_bg'], borderwidth=0)
        style.configure('TSpinbox', fieldbackground=DARK_THEME['entry_bg'], foreground=DARK_THEME['entry_fg'])
        style.configure('TPanedwindow', background=DARK_THEME['frame_bg'])
        root.configure(bg=DARK_THEME['bg'])

    @staticmethod
    def configure_text_widget(widget):
        widget.configure(bg=DARK_THEME['text_bg'], fg=DARK_THEME['text_fg'], insertbackground=DARK_THEME['insert_bg'], selectbackground=DARK_THEME['select_bg'], selectforeground=DARK_THEME['select_fg'], relief='flat', borderwidth=1, highlightthickness=1, highlightcolor=DARK_THEME['accent'], highlightbackground=DARK_THEME['border'])
        widget.tag_configure('error', foreground=DARK_THEME['error'], font=('Segoe UI', 9, 'bold'))
        widget.tag_configure('warning', foreground=DARK_THEME['warning'])
        widget.tag_configure('success', foreground=DARK_THEME['success'], font=('Segoe UI', 9, 'bold'))
        widget.tag_configure('info', foreground=DARK_THEME['accent'])

    @staticmethod
    def configure_listbox(widget):
        widget.configure(bg=DARK_THEME['text_bg'], fg=DARK_THEME['text_fg'], selectbackground=DARK_THEME['select_bg'], selectforeground=DARK_THEME['select_fg'], relief='flat', borderwidth=1, highlightthickness=1, highlightcolor=DARK_THEME['accent'], highlightbackground=DARK_THEME['border'])

class MarkdownPreviewWidget(ttk.Frame):
    def __init__(self, parent):
        super().__init__(parent)
        self.parent = parent
        self.markdown_processor = None
        if MARKDOWN_AVAILABLE: self.markdown_processor = markdown.Markdown(extensions=['codehilite', 'tables', 'fenced_code', 'toc'])
        self.create_preview_widgets()
        self.current_content = ""

    def create_preview_widgets(self):
        toolbar = ttk.Frame(self)
        toolbar.pack(fill=tk.X, pady=(0, 10), padx=5)
        ttk.Button(toolbar, text="🔄 Refresh", command=self.refresh_preview).pack(side=tk.LEFT, padx=(0, 5))
        ttk.Button(toolbar, text="🌐 Open in Browser", command=self.open_in_browser).pack(side=tk.LEFT, padx=(0, 5))
        ttk.Button(toolbar, text="💾 Export HTML", command=self.export_html).pack(side=tk.LEFT)
        self.preview_mode = tk.StringVar(value="rendered")
        ttk.Radiobutton(toolbar, text="Source", variable=self.preview_mode, value="source", command=self.switch_preview_mode).pack(side=tk.RIGHT)
        ttk.Radiobutton(toolbar, text="Rendered", variable=self.preview_mode, value="rendered", command=self.switch_preview_mode).pack(side=tk.RIGHT, padx=(0, 5))
        ttk.Label(toolbar, text="Preview:").pack(side=tk.RIGHT, padx=(10, 5))
        self.create_preview_area()
        self.show_placeholder()

    def create_preview_area(self):
        self.preview_text = scrolledtext.ScrolledText(self, wrap=tk.WORD, font=('Segoe UI', 10), state='disabled')
        self.preview_text.pack(fill=tk.BOTH, expand=True, padx=5, pady=(0, 5))
        DarkThemeManager.configure_text_widget(self.preview_text)
        self.configure_markdown_tags()

    def configure_markdown_tags(self):
        self.preview_text.tag_configure('h1', font=('Segoe UI', 18, 'bold'), foreground='#4CAF50', spacing1=12, spacing3=6)
        self.preview_text.tag_configure('h2', font=('Segoe UI', 16, 'bold'), foreground='#2196F3', spacing1=10, spacing3=5)
        self.preview_text.tag_configure('h3', font=('Segoe UI', 13, 'bold'), foreground='#FF9800', spacing1=8, spacing3=4)
        self.preview_text.tag_configure('h4', font=('Segoe UI', 11, 'bold'), foreground='#9C27B0', spacing1=6, spacing3=3)
        self.preview_text.tag_configure('bold', font=('Segoe UI', 10, 'bold'))
        self.preview_text.tag_configure('italic', font=('Segoe UI', 10, 'italic'))
        self.preview_text.tag_configure('code_inline', font=('Consolas', 9), background='#2d2d30', foreground='#dcdcaa')
        self.preview_text.tag_configure('code_block', font=('Consolas', 9), background='#1e1e1e', foreground='#d4d4d4', lmargin1=20, lmargin2=20, spacing1=5, spacing3=5, wrap='none')
        self.preview_text.tag_configure('list_item', lmargin1=20, lmargin2=40)
        self.preview_text.tag_configure('quote', lmargin1=20, lmargin2=20, background='#252526', foreground='#cccccc', spacing1=2, spacing3=2)
        self.preview_text.tag_configure('link', foreground='#4CAF50', underline=True)
        self.preview_text.tag_configure('separator', spacing1=10, spacing3=10, overstrike=True)

    def show_placeholder(self):
        placeholder = "# 📚 Study Guide Preview\n\nWelcome to the Enhanced Research & Study Guide Generator!\n\n## 🚀 Quick Start\n\n1.  **Choose your source**: Web scraping, local documents, or direct YouTube video analysis.\n2.  **Configure research**: Enable web and YouTube research (optional).\n3.  **Set up AI model**: Add your Gemini API key.\n4.  **Click \"Start Generation\"** to create your study guide.\n\nYour generated study guide will appear here with rich formatting!\n\n---\n\n*Ready when you are!*"
        self.update_preview(placeholder)

    def update_preview(self, content): self.current_content = content; self.refresh_preview()
    def refresh_preview(self):
        if not self.current_content: return
        self.preview_text.config(state='normal')
        self.preview_text.delete('1.0', tk.END)
        if self.preview_mode.get() == "source": self.preview_text.insert(tk.END, self.current_content)
        else: self.render_markdown_content(self.current_content)
        self.preview_text.config(state='disabled')

    def render_markdown_content(self, content):
        lines = content.split('\n'); i = 0
        while i < len(lines):
            line = lines[i].rstrip()
            if not line: self.preview_text.insert(tk.END, '\n'); i += 1; continue
            if line.startswith('# '): self.preview_text.insert(tk.END, line[2:] + '\n', 'h1')
            elif line.startswith('## '): self.preview_text.insert(tk.END, line[3:] + '\n', 'h2')
            elif line.startswith('### '): self.preview_text.insert(tk.END, line[4:] + '\n', 'h3')
            elif line.startswith('#### '): self.preview_text.insert(tk.END, line[5:] + '\n', 'h4')
            elif line.startswith('```'):
                self.preview_text.insert(tk.END, '\n'); i += 1; code_content = []
                while i < len(lines) and not lines[i].startswith('```'): code_content.append(lines[i]); i += 1
                if code_content: self.preview_text.insert(tk.END, '\n'.join(code_content) + '\n', 'code_block')
                self.preview_text.insert(tk.END, '\n')
            elif line.startswith('---') or line.startswith('==='): self.preview_text.insert(tk.END, ' ' * 80 + '\n', 'separator')
            elif line.startswith(('* ', '- ')) or re.match(r'^\d+\.\s', line): self.preview_text.insert(tk.END, '• ' + re.sub(r'^\* |^- |\d+\.\s', '', line) + '\n', 'list_item')
            elif line.startswith('> '): self.preview_text.insert(tk.END, line[2:] + '\n', 'quote')
            else: self.render_inline_formatting(line + '\n')
            i += 1

    def render_inline_formatting(self, text):
        parts = re.split(r'(\*\*.*?\*\*|\*.*?\*|`.*?`|\[.*?\]\(.*?\))', text)
        for part in parts:
            if part.startswith('**') and part.endswith('**'): self.preview_text.insert(tk.END, part[2:-2], 'bold')
            elif part.startswith('*') and part.endswith('*'): self.preview_text.insert(tk.END, part[1:-1], 'italic')
            elif part.startswith('`') and part.endswith('`'): self.preview_text.insert(tk.END, part[1:-1], 'code_inline')
            elif part.startswith('[') and '](' in part and part.endswith(')'): self.preview_text.insert(tk.END, part[1:part.index('](')], 'link')
            else: self.preview_text.insert(tk.END, part)

    def switch_preview_mode(self): self.refresh_preview()
    def open_in_browser(self):
        if not self.current_content: messagebox.showinfo("No Content", "No study guide content to display."); return
        try:
            html_content = self.generate_html_content(self.current_content)
            with tempfile.NamedTemporaryFile(mode='w', suffix='.html', delete=False, encoding='utf-8') as f: f.write(html_content); temp_path = f.name
            webbrowser.open(f'file://{os.path.abspath(temp_path)}')
            self.parent.after(5000, lambda: os.unlink(temp_path))
        except Exception as e: messagebox.showerror("Error", f"Failed to open in browser:\n{e}")

    def export_html(self):
        if not self.current_content: messagebox.showinfo("No Content", "No study guide content to export."); return
        initial_filename = f"study_guide_{datetime.now().strftime('%Y%m%d_%H%M')}.html"
        filepath = filedialog.asksaveasfilename(initialfile=initial_filename, defaultextension=".html", filetypes=[("HTML files", "*.html"), ("All files", "*.*")])
        if not filepath: return
        try:
            html_content = self.generate_html_content(self.current_content)
            with open(filepath, 'w', encoding='utf-8') as f: f.write(html_content)
            messagebox.showinfo("Success", f"HTML exported to:\n{filepath}")
        except Exception as e: messagebox.showerror("Error", f"Failed to export HTML:\n{e}")

    def generate_html_content(self, markdown_content):
        html_body = markdown.markdown(markdown_content, extensions=['fenced_code', 'tables', 'sane_lists']) if MARKDOWN_AVAILABLE else f"<pre>{markdown_content}</pre>"
        return f"""<!DOCTYPE html><html lang="en"><head><meta charset="UTF-8"><title>Study Guide</title><style>
body{{font-family:-apple-system,BlinkMacSystemFont,'Segoe UI',Roboto,Helvetica,Arial,sans-serif;line-height:1.6;max-width:800px;margin:2rem auto;padding:2rem;background-color:#1e1e1e;color:#d4d4d4}}
h1,h2,h3{{line-height:1.2;margin-top:1.5em}}h1{{color:#34a853}}h2{{color:#4285f4;border-bottom:1px solid #444}}h3{{color:#fbbc05}}
code{{background-color:#2d2d30;color:#dcdcaa;padding:.2em .4em;margin:0;font-size:.85em;border-radius:3px;font-family:monospace}}
pre{{background-color:#1e1e1e;border:1px solid #404040;border-radius:5px;padding:1em;overflow-x:auto}}pre code{{background:0 0;padding:0;font-size:1em}}
blockquote{{border-left:4px solid #34a853;margin:0;padding:0.5em 1em;color:#aaa}}a{{color:#4285f4}}hr{{border:none;height:1px;background-color:#404040;margin:2em 0}}
</style></head><body>{html_body}</body></html>"""

class APIKeyManager:
    @staticmethod
    def get_gemini_api_key(filepath: str) -> Optional[str]:
        if not Path(filepath).is_file(): logging.error(f"Gemini API key file not found at {filepath}"); return None
        try:
            with open(Path(filepath), 'r', encoding='utf-8') as f: key = f.read().strip()
            if not key: raise ValueError("Key file is empty.")
            logging.info(f"🔑 Gemini API key loaded successfully."); return key
        except Exception as e: logging.error(f"❌ Failed to load Gemini API key: {e}"); return None

    @staticmethod
    def get_google_search_credentials(key_path: str, cx_path: str) -> tuple[str | None, str | None]:
        try:
            with open(Path(key_path), 'r') as f: api_key = f.read().strip()
            with open(Path(cx_path), 'r') as f: cx_id = f.read().strip()
            if not api_key or not cx_id: raise ValueError("Key or CX is empty")
            logging.info("Google Search credentials loaded successfully.")
            return api_key, cx_id
        except Exception as e:
            logging.warning(f"Could not load Google Search credentials: {e}. Will use fallback.")
            return None, None

class LocalDocumentLoader:
    def __init__(self, file_paths: list[str]): self.file_paths = file_paths
    def _read_txt(self, p: Path) -> str: return p.read_text(encoding='utf-8', errors='ignore')
    def _read_pdf(self, p: Path) -> str:
        with fitz.open(p) as doc: return "".join(page.get_text() for page in doc)
    def _read_docx(self, p: Path) -> str:
        return "\n".join(para.text for para in docx.Document(p).paragraphs)
    def load_and_extract_text(self) -> dict[str, str]:
        content = {}
        for path_str in self.file_paths:
            p, name = Path(path_str), Path(path_str).name
            logging.info(f"Processing local file: {name}")
            try:
                if name.lower().endswith(".pdf"): content[name] = self._read_pdf(p)
                elif name.lower().endswith(".docx"): content[name] = self._read_docx(p)
                elif name.lower().endswith(".txt"): content[name] = self._read_txt(p)
                else: logging.warning(f"Unsupported file type for text-only extraction: {name}")
            except Exception as e: logging.error(f"Failed to process file '{name}': {e}")
        logging.info(f"Successfully processed {len(content)} local documents for text.")
        return content

class WebsiteScraper:
    def __init__(self, base_url: str, max_pages: int, user_agent: str, request_timeout: int, rate_limit_delay: float, gemini_api_key: Optional[str] = None, deep_crawl: bool = False, crawl_depth: int = 0):
        self.base_url = base_url
        self.domain = urlparse(base_url).netloc
        self.max_pages = max_pages
        self.session = requests.Session()
        self.session.headers.update({'User-Agent': user_agent})
        self.request_timeout = request_timeout
        self.rate_limit_delay = rate_limit_delay
        self.scraped_content = {}
        self.visited_urls = set()
        self.gemini_api_key = gemini_api_key
        self.llm = None
        self.crawl_topic = ""
        self.deep_crawl = deep_crawl
        self.crawl_depth = crawl_depth if crawl_depth > 0 else float('inf')

        if self.gemini_api_key and DEPENDENCIES_AVAILABLE:
            try:
                self.llm = ChatGoogleGenerativeAI(model="gemini-1.5-flash", google_api_key=self.gemini_api_key, temperature=0.0)
                logging.info("🧠 WebsiteScraper initialized with LLM for intelligent link selection.")
            except Exception as e:
                logging.warning(f"⚠️ Could not initialize LLM for scraper, will use basic selection. Error: {e}")
                self.llm = None

    def _get_crawl_topic(self, initial_content: str) -> str:
        if not self.llm: return "general information"
        try:
            prompt = ChatPromptTemplate.from_template("Summarize the main topic of the following text in a short, concise phrase. Examples: 'asynchronous programming in Python', 'React state management libraries'.\n\nTEXT CONTENT:\n---\n{content}\n---\n\nMain Topic:")
            chain = prompt | self.llm | StrOutputParser()
            topic = chain.invoke({"content": initial_content})
            return topic.strip()
        except Exception as e:
            logging.error(f"❌ Could not determine crawl topic with LLM: {e}")
            return "general information"

    def _select_relevant_urls_with_llm(self, links_with_text: List[Tuple[str, str]], limit: int) -> List[str]:
        if not self.llm or not self.crawl_topic: return [link[0] for link in links_with_text[:limit]]
        formatted_links = "\n".join([f"- URL: {url}\n  Link Text: '{text}'" for url, text in links_with_text])
        try:
            prompt = ChatPromptTemplate.from_template(
                "You are an intelligent web crawler. Based on the **Main Topic**, select the **{limit}** most relevant URLs from the list. Prioritize tutorials, guides, and core features. Avoid 'contact', 'about', 'pricing', 'login'.\n\n"
                "**Main Topic:** {topic}\n\n"
                "**Candidate Links:**\n{links}\n\n"
                "Respond ONLY with a JSON list of URL strings. Example: [\"https://.../url1\", \"https://.../url2\"]"
            )
            chain = prompt | self.llm | StrOutputParser()
            response = chain.invoke({"topic": self.crawl_topic, "links": formatted_links, "limit": limit})
            json_match = re.search(r'\[.*\]', response, re.DOTALL)
            if json_match:
                selected_urls = json.loads(json_match.group(0))
                if isinstance(selected_urls, list) and all(isinstance(i, str) for i in selected_urls):
                    return selected_urls[:limit]
            logging.warning("⚠️ LLM response for URL selection was not valid JSON. Falling back.")
        except Exception as e:
            logging.error(f"❌ LLM URL selection failed: {e}. Falling back.")
        return [link[0] for link in links_with_text[:limit]]

    def _get_page_content(self, url: str) -> str | None:
        try:
            response = self.session.get(url, timeout=self.request_timeout)
            response.raise_for_status()
            if 'text/html' in response.headers.get('Content-Type', ''): return response.text
        except requests.exceptions.RequestException as e: logging.warning(f"Request failed for {url}: {e}")
        return None

    def _parse_content_and_links(self, html: str, page_url: str) -> tuple[str, list[tuple[str, str]]]:
        soup = BeautifulSoup(html, 'html.parser')
        for element in soup.select('script, style, nav, footer, header, aside, form, .ad, .advertisement'): element.decompose()
        main_content_area = (soup.select_one('main, article, [role="main"], .main-content, .content, .post-content') or soup.body)
        text_maker = html2text.HTML2Text(); text_maker.body_width = 0; text_maker.ignore_links = True; text_maker.ignore_images = True
        markdown_content = text_maker.handle(str(main_content_area)) if main_content_area else ""
        markdown_content = re.sub(r'\n{3,}', '\n\n', markdown_content).strip()
        markdown_content = re.sub(r'[ \t]+', ' ', markdown_content)
        links = set()
        for a in soup.find_all('a', href=True):
            href = a['href']
            if href.startswith(('#', 'mailto:', 'tel:')): continue
            full_url = urljoin(page_url, href).split('#')[0]
            if urlparse(full_url).scheme in ['http', 'https'] and urlparse(full_url).netloc == self.domain:
                anchor_text = a.get_text(strip=True)
                if anchor_text and len(anchor_text) > 2:
                    links.add((full_url, anchor_text))
        return markdown_content, list(links)

    def crawl(self, additional_urls: Optional[List[str]] = None) -> dict[str, str]:
        # Check for YouTube URL in normal URL mode
        if not additional_urls and ('youtube.com/watch' in self.base_url or 'youtu.be/' in self.base_url):
            logging.info("▶️ YouTube URL detected in web scraper mode. Switching to subtitle download.")
            if not YT_DLP_AVAILABLE:
                logging.error("❌ yt-dlp is not available. Cannot download YouTube subtitles.")
                return {}

            video_id_match = re.search(r'(?:v=|\/|embed\/|watch\?v=)([a-zA-Z0-9_-]{11})', self.base_url)
            if video_id_match:
                video_id = video_id_match.group(1)
                logging.info(f"Found YouTube video ID: {video_id}. Attempting to download transcript.")
                yt_researcher = EnhancedYouTubeResearcher()
                try:
                    transcript = yt_researcher.download_transcript(video_id)
                    if transcript and len(transcript.strip()) > 50:
                        logging.info(f"✅ Successfully downloaded transcript for video ID {video_id}.")
                        # To keep the output consistent, let's add some metadata like the other YouTube method does
                        metadata = yt_researcher._get_video_metadata(video_id)
                        title = metadata.get('title', 'YouTube Video') if metadata else 'YouTube Video'
                        full_content = f"Video Title: {title}\nURL: {self.base_url}\n\nTranscript:\n{transcript}"
                        self.scraped_content[self.base_url] = full_content
                    else:
                        logging.warning(f"⚠️ Could not retrieve a valid transcript for video ID {video_id}.")
                except Exception as e:
                    logging.error(f"❌ An error occurred during YouTube transcript download: {e}")
                finally:
                    yt_researcher._cleanup() # Ensure temporary files are removed
                return self.scraped_content
            else:
                logging.warning(f"Could not parse a valid video ID from URL: {self.base_url}")
                return {}

        if additional_urls:
            logging.info(f"Starting flat crawl of {len(additional_urls)} external research URLs...")
            for url in additional_urls:
                if len(self.scraped_content) >= self.max_pages: break
                if url in self.visited_urls: continue
                self.visited_urls.add(url)
                logging.info(f"[{len(self.scraped_content) + 1}/{self.max_pages}] Research page: {url}")
                html = self._get_page_content(url)
                if html:
                    text, _ = self._parse_content_and_links(html, url)
                    if text and len(text.strip()) > 150: self.scraped_content[url] = text
                sleep(self.rate_limit_delay)
            logging.info(f"Research crawling complete. Scraped {len(self.scraped_content)} pages.")
            return self.scraped_content

        if self.deep_crawl:
             logging.info(f"🚀 Starting DEEP crawl from {self.base_url} (Depth: {self.crawl_depth if self.crawl_depth != float('inf') else 'Infinite'})")
        else:
             logging.info(f"🚀 Starting ONE-LEVEL crawl from {self.base_url}")

        urls_to_visit = collections.deque([(self.base_url, 0)])

        while urls_to_visit and len(self.scraped_content) < self.max_pages:
            url, current_depth = urls_to_visit.popleft()
            if url in self.visited_urls: continue
            self.visited_urls.add(url)

            logging.info(f"[{len(self.scraped_content) + 1}/{self.max_pages}][Depth:{current_depth}] Scraping: {url}")
            html = self._get_page_content(url)
            if not html: continue

            text, new_links_with_text = self._parse_content_and_links(html, url)
            if text and len(text.strip()) > 150: self.scraped_content[url] = text
            if not self.crawl_topic and self.llm and text:
                self.crawl_topic = self._get_crawl_topic(text[:4000])
                logging.info(f"🧠 Determined crawl context: '{self.crawl_topic}'")

            should_add_links = self.deep_crawl or current_depth == 0
            if current_depth >= self.crawl_depth: should_add_links = False

            if should_add_links and new_links_with_text:
                remaining_capacity = self.max_pages - len(self.scraped_content)
                unvisited_links = [(u, t) for u, t in new_links_with_text if u not in self.visited_urls and u not in [item[0] for item in urls_to_visit]]
                if not unvisited_links: continue
                if self.llm and self.crawl_topic and len(unvisited_links) > remaining_capacity:
                    logging.info(f"🧠 Found {len(unvisited_links)} links, exceeds capacity. Using LLM to select...")
                    selected_urls = self._select_relevant_urls_with_llm(unvisited_links, limit=remaining_capacity)
                    logging.info(f"  ✅ LLM selected {len(selected_urls)} URLs to visit next.")
                else:
                    selected_urls = [u for u, t in unvisited_links[:remaining_capacity]]
                for new_url in selected_urls: urls_to_visit.append((new_url, current_depth + 1))
            sleep(self.rate_limit_delay)
        logging.info(f"Crawl complete. Scraped {len(self.scraped_content)} total pages.")
        return self.scraped_content

class EnhancedResearchQueryGenerator:
    def __init__(self, api_key: str):
        self.llm = ChatGoogleGenerativeAI(model="gemini-1.5-flash", google_api_key=api_key, temperature=0.3)
        self.output_parser = StrOutputParser()

    def extract_main_topic_and_subtopics(self, content: dict[str, str]) -> Dict[str, List[str]]:
        logging.info("🧠 Analyzing content structure and extracting topics...")
        context = ""
        for name, text in list(content.items())[:3]:
            truncated_text = text[:2000] if len(text) > 2000 else text
            context += f"\n--- {name} ---\n{truncated_text}"
        if len(context) > 6000: context = context[:6000] + "..."
        prompt = ChatPromptTemplate.from_template('Analyze this content and extract:\n1. Main topic (the primary subject/technology)\n2. Key subtopics (3-5 related concepts, features, or areas)\n\nReturn as JSON:\n{{"main_topic": "topic name", "subtopics": ["subtopic1", "subtopic2", ...]}}\n\nCONTENT:\n{context}\n\nJSON Response:')
        try:
            response = (prompt | self.llm | self.output_parser).invoke({"context": context})
            json_match = re.search(r'\{.*\}', response, re.DOTALL)
            if json_match:
                result = json.loads(json_match.group(0))
                main_topic = result.get('main_topic', 'General Topic')
                subtopics = result.get('subtopics', [])
                logging.info(f"✅ Main topic: {main_topic}"); logging.info(f"✅ Subtopics: {', '.join(subtopics)}")
                return {"main_topic": main_topic, "subtopics": subtopics}
        except Exception as e: logging.error(f"❌ Topic extraction failed: {e}")
        return {"main_topic": "General Topic", "subtopics": []}

    def generate_diverse_research_queries(self, topic_data: Dict[str, List[str]], max_queries: int = 8) -> List[str]:
        logging.info(f"🧠 Generating diverse research queries...")
        main_topic, subtopics = topic_data["main_topic"], topic_data["subtopics"]
        query_templates = [f"{main_topic} complete tutorial guide", f"{main_topic} best practices production", f"{main_topic} common errors troubleshooting", f"{main_topic} advanced techniques tips", f"{main_topic} performance optimization", f"{main_topic} vs alternatives comparison", f"how to use {main_topic} examples", f"{main_topic} latest updates 2024"]
        for subtopic in subtopics[:3]: query_templates.extend([f"{main_topic} {subtopic} tutorial", f"{subtopic} {main_topic} implementation"])
        prompt = ChatPromptTemplate.from_template('Based on the topic "{main_topic}" and subtopics {subtopics}, generate {max_queries} diverse search queries.\nInclude queries for: tutorials, troubleshooting, best practices, comparisons, and recent updates.\n\nReturn as a JSON array of strings only:')
        try:
            response = (prompt | self.llm | self.output_parser).invoke({"main_topic": main_topic, "subtopics": subtopics, "max_queries": max_queries})
            json_match = re.search(r'\[.*\]', response, re.DOTALL)
            if json_match:
                llm_queries = json.loads(json_match.group(0))
                if isinstance(llm_queries, list) and len(llm_queries) > 0:
                    logging.info(f"✅ Generated {len(llm_queries)} LLM queries"); return llm_queries[:max_queries]
        except Exception as e: logging.warning(f"⚠️ LLM query generation failed, using templates: {e}")
        selected_queries = query_templates[:max_queries]; logging.info(f"✅ Using {len(selected_queries)} template queries"); return selected_queries

class EnhancedGoogleSearchResearcher:
    def __init__(self, api_key: str = None, cx_id: str = None):
        self.api_key, self.cx_id = api_key, cx_id
        self.session = requests.Session(); self.session.headers.update({'User-Agent': random.choice(USER_AGENTS)})
    def _is_quality_url(self, url: str, exclude_domain: str) -> bool:
        if not url or (exclude_domain and exclude_domain in url): return False
        bad_extensions = ['.pdf', '.zip', '.doc', '.docx', '.ppt', '.pptx']
        low_quality_domains = ['pinterest.com', 'facebook.com', 'twitter.com', 'instagram.com']
        if any(url.lower().endswith(ext) for ext in bad_extensions): return False
        parsed_url = urlparse(url)
        if any(domain in parsed_url.netloc.lower() for domain in low_quality_domains): return False
        return True
    def search_and_extract_urls(self, queries: list[str], exclude_domain: str, max_results_per_query: int = 8) -> list[str]:
        all_urls, successful_queries = set(), 0
        for i, query in enumerate(queries):
            logging.info(f"🔍 Searching [{i+1}/{len(queries)}]: {query}")
            try:
                if self.api_key and self.cx_id:
                    params = {'key': self.api_key, 'cx': self.cx_id, 'q': query, 'num': max_results_per_query, 'dateRestrict': 'y2'}
                    response = self.session.get("https://www.googleapis.com/customsearch/v1", params=params, timeout=10)
                    response.raise_for_status()
                    results = response.json().get('items', [])
                    query_urls = [item['link'] for item in results]
                else:
                    params = {'q': query, 'kl': 'us-en'}
                    response = self.session.get("https://html.duckduckgo.com/html/", params=params, timeout=15)
                    response.raise_for_status()
                    soup = BeautifulSoup(response.text, 'html.parser'); query_urls = []
                    for link in soup.select('a.result__a'):
                        href = link.get('href', '')
                        if 'uddg=' in href:
                            try: query_urls.append(unquote(href.split('uddg=')[1]))
                            except: continue
                quality_urls = [url for url in query_urls if self._is_quality_url(url, exclude_domain)]
                all_urls.update(quality_urls); successful_queries += 1
                logging.info(f"  ✅ Found {len(quality_urls)} quality URLs"); sleep(random.uniform(1.0, 2.0))
            except Exception as e: logging.error(f"  ❌ Search failed for '{query}': {e}")
        logging.info(f"🔍 Search complete: {len(all_urls)} unique URLs from {successful_queries}/{len(queries)} queries")
        return list(all_urls)

class EnhancedPlaywrightResearcher:
    def __init__(self):
        self.search_engines = [{"name": "DuckDuckGo", "url": "https://duckduckgo.com/", "input_selector": 'input[name="q"]'}, {"name": "Bing", "url": "https://www.bing.com/", "input_selector": 'input[name="q"]'}]
        logging.info("🤖 Initialized Enhanced Playwright Researcher")
    def search_and_extract_urls(self, queries: List[str], exclude_domain: str) -> List[str]:
        all_urls = set()
        with sync_playwright() as p:
            browser = p.chromium.launch(headless=True)
            context = browser.new_context(user_agent=random.choice(USER_AGENTS)); page = context.new_page()
            try:
                for i, query in enumerate(queries):
                    logging.info(f"🤖 Playwright search [{i+1}/{len(queries)}]: {query}")
                    for engine in self.search_engines:
                        try:
                            page.goto(engine["url"], timeout=20000)
                            page.locator(engine["input_selector"]).fill(query)
                            page.locator(engine["input_selector"]).press("Enter")
                            result_selectors = ["div#links", ".b_results", "[data-testid='result']"]
                            for selector in result_selectors:
                                try: page.wait_for_selector(selector, timeout=10000); break
                                except: continue
                            link_selectors = ["h2 > a", ".b_title > h2 > a", "[data-testid='result-title-a']"]; found_urls = set()
                            for selector in link_selectors:
                                try:
                                    links = page.locator(selector).all()
                                    for link in links:
                                        href = link.get_attribute('href')
                                        if href and self._is_useful_url(href, exclude_domain): found_urls.add(href)
                                except: continue
                            all_urls.update(found_urls); logging.info(f"  ✅ {engine['name']}: {len(found_urls)} URLs"); sleep(random.uniform(2.0, 3.0)); break
                        except Exception as e: logging.warning(f"  ⚠️ {engine['name']} failed: {e}"); continue
                    sleep(random.uniform(3.0, 5.0))
            finally: browser.close()
        logging.info(f"🤖 Playwright search complete: {len(all_urls)} unique URLs")
        return list(all_urls)
    def _is_useful_url(self, url: str, domain: str) -> bool:
        if not url or (domain and domain in url): return False
        bad_extensions = ['.pdf', '.zip', '.doc', '.docx', '.ppt', '.pptx', '.exe']
        bad_domains = ['pinterest.com', 'facebook.com', 'twitter.com', 'instagram.com', 'tiktok.com']
        url_lower = url.lower()
        if any(url_lower.endswith(ext) for ext in bad_extensions): return False
        parsed_url = urlparse(url)
        if any(bad_domain in parsed_url.netloc.lower() for bad_domain in bad_domains): return False
        return True

class EnhancedYouTubeResearcher:
    def __init__(self):
        if not YT_DLP_AVAILABLE: raise ImportError("YouTube research requires 'yt-dlp'. Please install it.")
        self.temp_dirs = []; logging.info("🔎 Initialized Enhanced YouTube Researcher")
    def _clean_transcript_text(self, vtt_content: str) -> str:
        if not vtt_content: return ""
        lines = vtt_content.splitlines(); text_lines = []
        for line in lines:
            line = line.strip()
            if (not line or line.startswith('WEBVTT') or '-->' in line or line.startswith('NOTE') or re.match(r'^\d+$', line)): continue
            line = re.sub(r'<[^>]+>', '', line); line = re.sub(r'&[a-zA-Z]+;', '', line); text_lines.append(line)
        full_text = ' '.join(text_lines)
        full_text = re.sub(r'\[Music\]|\[Applause\]|\[Laughter\]', '', full_text); full_text = re.sub(r'\s+', ' ', full_text)
        return full_text.strip()
    def _extract_transcript_from_file(self, file_path: Path) -> str:
        try:
            if not file_path.exists(): return ""
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f: content = f.read()
            return self._clean_transcript_text(content)
        except Exception as e: logging.warning(f"Failed to read transcript file {file_path}: {e}"); return ""
    def _get_video_metadata(self, video_id: str) -> Optional[Dict]:
        try:
            with yt_dlp.YoutubeDL({'quiet': True, 'no_warnings': True, 'extract_flat': False}) as ydl:
                info = ydl.extract_info(f"https://www.youtube.com/watch?v={video_id}", download=False)
                return {'id': info.get('id'), 'title': info.get('title'), 'duration': info.get('duration', 0), 'view_count': info.get('view_count', 0), 'upload_date': info.get('upload_date'), 'subtitles': info.get('subtitles', {}), 'automatic_captions': info.get('automatic_captions', {})}
        except Exception: return None
    def _has_quality_subtitles(self, metadata: Dict) -> bool:
        if not metadata: return False
        subtitles, auto_captions = metadata.get('subtitles', {}), metadata.get('automatic_captions', {})
        english_subs = ['en', 'en-US', 'en-GB', 'en-CA', 'en-AU']
        for lang in english_subs:
            if lang in subtitles: return True
        for lang in english_subs:
            if lang in auto_captions: return True
        return False
    def download_transcript(self, video_id: str, title: str = "") -> Optional[str]:
        video_url = f"https://www.youtube.com/watch?v={video_id}"; temp_dir = Path(f"./temp_subs_{video_id}_{random.randint(1000, 9999)}"); temp_dir.mkdir(exist_ok=True); self.temp_dirs.append(temp_dir)
        try:
            ydl_opts = {'writesubtitles': True, 'writeautomaticsub': True, 'subtitleslangs': ['en', 'en-US', 'en-GB', 'en-CA', 'en-AU'], 'skip_download': True, 'outtmpl': str(temp_dir / '%(id)s'), 'quiet': True, 'no_warnings': True}
            with yt_dlp.YoutubeDL(ydl_opts) as ydl: ydl.download([video_url])
            subtitle_files = list(temp_dir.glob('*.vtt'))
            if not subtitle_files: return None
            manual_subs = [f for f in subtitle_files if '.auto.' not in f.name]; subtitle_file = manual_subs[0] if manual_subs else subtitle_files[0]
            transcript = self._extract_transcript_from_file(subtitle_file)
            if len(transcript) < 100: return None
            return transcript
        except Exception as e: logging.warning(f"Failed to download transcript for {video_id}: {e}"); return None
    def search_videos_with_yt_dlp(self, query: str, max_results: int = 10) -> List[Dict]:
        try:
            with yt_dlp.YoutubeDL({'quiet': True, 'no_warnings': True, 'extract_flat': True}) as ydl: search_results = ydl.extract_info(f"ytsearch{max_results*2}:{query}", download=False)
            if not search_results or 'entries' not in search_results: return []
            videos = []
            for entry in search_results['entries']:
                if not entry or not entry.get('id'): continue
                metadata = self._get_video_metadata(entry['id'])
                if not metadata: continue
                duration, view_count = metadata.get('duration', 0), metadata.get('view_count', 0)
                if duration < 120 or duration > 3600: continue
                if view_count < 1000: continue
                if not self._has_quality_subtitles(metadata): continue
                videos.append(metadata)
                if len(videos) >= max_results: break
            return videos
        except Exception as e: logging.error(f"Video search failed for '{query}': {e}"); return []
    def get_transcripts_for_queries(self, queries: List[str], max_videos_per_query: int = 3) -> Dict[str, str]:
        logging.info(f"▶️  Starting enhanced YouTube research for {len(queries)} queries..."); all_videos = {}
        for i, query in enumerate(queries):
            logging.info(f"🔍 YouTube search [{i+1}/{len(queries)}]: {query}")
            videos = self.search_videos_with_yt_dlp(query, max_videos_per_query)
            for video in videos:
                if video['id'] not in all_videos: all_videos[video['id']] = video
            logging.info(f"  ✅ Found {len(videos)} quality videos"); sleep(1.0)
        if not all_videos: logging.warning("No suitable videos found for transcript extraction"); return {}
        logging.info(f"📺 Found {len(all_videos)} unique videos. Downloading transcripts..."); transcripts = {}
        def download_single_transcript(video_data):
            video_id, metadata = video_data; transcript = self.download_transcript(video_id, metadata.get('title', ''))
            if transcript: return video_id, metadata, transcript
            return None
        with ThreadPoolExecutor(max_workers=3) as executor:
            future_to_video = {executor.submit(download_single_transcript, item): item[0] for item in all_videos.items()}
            for future in as_completed(future_to_video):
                try:
                    result = future.result(timeout=60)
                    if result:
                        video_id, metadata, transcript = result; title = metadata.get('title', 'Unknown Title'); duration = metadata.get('duration', 0); view_count = metadata.get('view_count', 0)
                        content = f"Video Title: {title}\nDuration: {duration//60}:{duration%60:02d}\nViews: {view_count:,}\nURL: https://www.youtube.com/watch?v={video_id}\n\nTranscript:\n{transcript}"
                        transcripts[f"https://www.youtube.com/watch?v={video_id}"] = content; logging.info(f"  ✅ Extracted transcript: {title[:50]}...")
                except Exception as e: logging.error(f"  ❌ Error processing video: {e}")
        self._cleanup()
        logging.info(f"▶️  YouTube research complete: {len(transcripts)} transcripts extracted"); return transcripts
    def _cleanup(self):
        for temp_dir in self.temp_dirs:
            try:
                if temp_dir.exists(): shutil.rmtree(temp_dir)
            except Exception as e: logging.warning(f"Failed to cleanup temp dir {temp_dir}: {e}")
        self.temp_dirs.clear()

class EnhancedNoteGenerator:
    def __init__(self, api_key: str, llm_config: dict, prompt_template_string: str):
        self.llm = ChatGoogleGenerativeAI(model=llm_config['model_name'], google_api_key=api_key, **llm_config['parameters'])
        self.output_parser = StrOutputParser(); self.prompt = ChatPromptTemplate.from_template(prompt_template_string)
    def _prepare_content_for_generation(self, source_data: dict[str, str]) -> str:
        if not source_data: return ""
        web_content, video_content, local_content = [], [], []
        for source_url, content in source_data.items():
            if 'youtube.com' in source_url: video_content.append((source_url, content))
            elif source_url.startswith('http'): web_content.append((source_url, content))
            else: local_content.append((source_url, content))
        organized_content = ""
        if local_content:
            organized_content += "\n\n=== LOCAL DOCUMENTS ===\n"
            for name, content in local_content: organized_content += f"\n--- SOURCE: {name} ---\n{content[:3000]}\n"
        if web_content:
            organized_content += "\n\n=== WEB RESEARCH ===\n"
            for url, content in web_content[:5]: organized_content += f"\n--- SOURCE: {url} ---\n{content[:2000]}\n"
        if video_content:
            organized_content += "\n\n=== VIDEO TRANSCRIPTS ===\n"
            for url, content in video_content: organized_content += f"\n--- SOURCE: {url} ---\n{content[:2500]}\n"
        return organized_content
    def generate_comprehensive_notes(self, source_data: dict[str, str], source_name: str) -> str:
        if not source_data: return "No content was provided to generate notes."
        logging.info(f"📝 Generating comprehensive notes from {len(source_data)} sources...")
        organized_content = self._prepare_content_for_generation(source_data)
        metadata = f"GENERATION METADATA:\n- Total Sources: {len(source_data)}\n- Source Types: {self._get_source_type_summary(source_data)}\n- Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n- Primary Source: {source_name}\n\n"
        try:
            chain = self.prompt | self.llm | self.output_parser
            notes = chain.invoke({"content": organized_content, "website_url": source_name, "source_count": len(source_data), "metadata": metadata})
            final_notes = metadata + "="*80 + "\n\n" + notes; logging.info("✅ Note generation completed successfully"); return final_notes
        except Exception as e:
            logging.error(f"❌ Error during note generation: {e}")
            return f"# Generation Error\n\nAn error occurred while communicating with the AI model:\n\n**Error Details:** `{e}`\n\n**Available Content Summary:**\n- Total sources processed: {len(source_data)}\n- Content length: {len(organized_content)} characters\n\nPlease check your API key and try again."
    def _get_source_type_summary(self, source_data: dict[str, str]) -> str:
        types = {"Web": 0, "YouTube": 0, "Local": 0}
        for source_url in source_data.keys():
            if 'youtube.com' in source_url: types["YouTube"] += 1
            elif source_url.startswith('http'): types["Web"] += 1
            else: types["Local"] += 1
        return ", ".join([f"{k}: {v}" for k, v in types.items() if v > 0])

class QueueHandler(logging.Handler):
    def __init__(self, log_queue): super().__init__(); self.log_queue = log_queue
    def emit(self, record): self.log_queue.put(f"[{datetime.now().strftime('%H:%M:%S')}] {self.format(record)}")

class AdvancedScraperApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Enhanced Research & Study Guide Generator v2.4 (Multimodal)")
        self.root.geometry("1400x900")
        self.root.minsize(1100, 700)
        DarkThemeManager.configure_dark_theme(self.root)
        self.input_mode_var = tk.StringVar(value="scrape")
        self.url_var = tk.StringVar()
        self.limit_var = tk.IntVar(value=10)
        self.deep_crawl_var = tk.BooleanVar(value=False)
        self.crawl_depth_var = tk.IntVar(value=2)
        self.research_enabled_var = tk.BooleanVar(value=False)
        self.web_research_enabled_var = tk.BooleanVar(value=True)
        self.yt_research_enabled_var = tk.BooleanVar(value=True)
        self.web_research_method_var = tk.StringVar(value="google_api")
        self.research_pages_var = tk.IntVar(value=5)
        self.research_queries_var = tk.IntVar(value=6)
        self.google_api_key_file_var = tk.StringVar()
        self.google_cx_file_var = tk.StringVar()
        self.yt_videos_per_query_var = tk.IntVar(value=3)
        self.api_key_file_var = tk.StringVar()
        self.model_name_var = tk.StringVar()
        self.temperature_var = tk.DoubleVar()
        self.max_tokens_var = tk.IntVar()
        ### NEW ###
        self.multimodal_upload_var = tk.BooleanVar(value=False)
        self.final_notes_content = ""
        self.config = {}
        self.is_processing = False
        self.create_widgets()
        self.load_initial_settings()
        self.toggle_input_mode()
        self.toggle_research_panel()
        self.setup_logging()

    def create_widgets(self):
        main_paned = ttk.PanedWindow(self.root, orient=tk.HORIZONTAL)
        main_paned.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        settings_frame = ttk.Frame(main_paned, width=420)
        main_paned.add(settings_frame, weight=0)
        settings_notebook = ttk.Notebook(settings_frame)
        settings_notebook.pack(fill="both", expand=True, pady=(0, 10))
        source_tab, research_tab, ai_tab, prompt_tab = (ttk.Frame(settings_notebook, padding=15) for _ in range(4))
        self.create_source_tab(source_tab)
        self.create_research_tab(research_tab)
        self.create_ai_tab(ai_tab)
        self.create_prompt_tab(prompt_tab)
        settings_notebook.add(source_tab, text="📄 Source")
        settings_notebook.add(research_tab, text="🔍 Research")
        settings_notebook.add(ai_tab, text="🤖 AI Model")
        settings_notebook.add(prompt_tab, text="📝 Prompt")
        button_frame = ttk.Frame(settings_frame)
        button_frame.pack(fill=tk.X, pady=(10, 0))
        self.start_button = ttk.Button(button_frame, text="🚀 Start Generation", command=self.start_task_thread, style="Accent.TButton")
        self.start_button.pack(fill=tk.X, ipady=5, pady=(0, 5))
        self.save_button = ttk.Button(button_frame, text="💾 Save Study Guide", command=self.save_notes, state=tk.DISABLED)
        self.save_button.pack(fill=tk.X, ipady=2)
        self.progress_var = tk.StringVar(value="Ready")
        ttk.Label(button_frame, textvariable=self.progress_var, foreground=DARK_THEME['accent'], anchor='center').pack(pady=(5, 0), fill=tk.X)
        content_notebook = ttk.Notebook(main_paned)
        main_paned.add(content_notebook, weight=1)
        self.markdown_preview = MarkdownPreviewWidget(content_notebook)
        self.log_text_widget = scrolledtext.ScrolledText(content_notebook, state='disabled', wrap=tk.WORD, font=("Consolas", 9))
        DarkThemeManager.configure_text_widget(self.log_text_widget)
        content_notebook.add(self.markdown_preview, text="📖 Study Guide Preview")
        content_notebook.add(self.log_text_widget, text="📋 Process Logs")

    def create_source_tab(self, parent):
        mode_frame = ttk.LabelFrame(parent, text="📥 Input Mode", padding=10)
        mode_frame.pack(fill=tk.X, pady=(0, 10))
        ttk.Radiobutton(mode_frame, text="🌐 Web Scraper (Intelligent Crawler)", variable=self.input_mode_var, value="scrape", command=self.toggle_input_mode).pack(anchor=tk.W)
        ttk.Radiobutton(mode_frame, text="🎬 YouTube Video (URL Analysis)", variable=self.input_mode_var, value="youtube_video", command=self.toggle_input_mode).pack(anchor=tk.W)
        ttk.Radiobutton(mode_frame, text="📁 Local Files (Text or Multimodal)", variable=self.input_mode_var, value="upload", command=self.toggle_input_mode).pack(anchor=tk.W)

        self.url_input_frame = ttk.LabelFrame(parent, text="🌐 URL Input", padding=10)
        self.url_input_frame.pack(fill=tk.X, pady=(0, 10))
        ttk.Label(self.url_input_frame, text="Target URL:").pack(fill=tk.X, anchor='w')
        ttk.Entry(self.url_input_frame, textvariable=self.url_var, font=("Segoe UI", 9)).pack(fill=tk.X, pady=(2, 8))
        self.scraper_options_frame = ttk.Frame(self.url_input_frame)
        self.scraper_options_frame.pack(fill=tk.X, pady=(0, 0))
        ttk.Label(self.scraper_options_frame, text="Max total pages to scrape:").pack(fill=tk.X, anchor='w')
        ttk.Spinbox(self.scraper_options_frame, from_=1, to=1000, textvariable=self.limit_var, width=10).pack(fill=tk.X, pady=(2, 10))
        self.deep_crawl_check = ttk.Checkbutton(self.scraper_options_frame, text="Enable Deep Crawl (Follows links recursively)", variable=self.deep_crawl_var, command=self.toggle_depth_setting)
        self.deep_crawl_check.pack(anchor=tk.W)
        self.depth_frame = ttk.Frame(self.scraper_options_frame)
        self.depth_frame.pack(fill=tk.X, padx=(20, 0), pady=(2, 0))
        self.depth_label = ttk.Label(self.depth_frame, text="Crawl Depth (0 for infinite):")
        self.depth_label.pack(side=tk.LEFT, padx=(0, 5))
        self.depth_spinbox = ttk.Spinbox(self.depth_frame, from_=0, to=50, textvariable=self.crawl_depth_var, width=8)
        self.depth_spinbox.pack(side=tk.LEFT)
        self.upload_frame = ttk.LabelFrame(parent, text="📂 Local File Settings", padding=10)
        self.upload_frame.pack(fill=tk.X)
        ### NEW ###
        self.multimodal_check = ttk.Checkbutton(self.upload_frame, text="Enable Multimodal Analysis (for Videos, Images, etc.)", variable=self.multimodal_upload_var)
        self.multimodal_check.pack(anchor=tk.W, pady=(0, 10))

        btn_frame = ttk.Frame(self.upload_frame)
        btn_frame.pack(fill=tk.X, pady=(0, 5))
        ttk.Button(btn_frame, text="➕ Add Files", command=self.add_files).pack(side=tk.LEFT, expand=True, fill=tk.X, padx=(0, 5))
        ttk.Button(btn_frame, text="🗑️ Clear List", command=self.clear_files).pack(side=tk.RIGHT, expand=True, fill=tk.X)
        listbox_frame = ttk.Frame(self.upload_frame); listbox_frame.pack(fill=tk.BOTH, expand=True)
        self.file_listbox = tk.Listbox(listbox_frame, height=5, font=("Segoe UI", 8))
        scrollbar = ttk.Scrollbar(listbox_frame, orient=tk.VERTICAL, command=self.file_listbox.yview)
        self.file_listbox.configure(yscrollcommand=scrollbar.set)
        DarkThemeManager.configure_listbox(self.file_listbox)
        self.file_listbox.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)

    def create_research_tab(self, parent):
        master_frame = ttk.LabelFrame(parent, text="🎛️ Master Control", padding=10)
        master_frame.pack(fill=tk.X, pady=(0, 10))
        self.research_checkbutton = ttk.Checkbutton(master_frame, text="🔬 Enable AI-Powered Research (Beta)", variable=self.research_enabled_var, command=self.toggle_research_panel)
        self.research_checkbutton.pack(anchor=tk.W)
        settings_frame = ttk.LabelFrame(parent, text="⚙️ Research Settings", padding=10)
        settings_frame.pack(fill=tk.X, pady=(0, 10))
        ttk.Label(settings_frame, text="AI search queries to generate:").pack(anchor=tk.W)
        ttk.Spinbox(settings_frame, from_=3, to=15, textvariable=self.research_queries_var, width=10).pack(fill=tk.X, pady=(2, 10))
        self.web_research_panel = ttk.LabelFrame(parent, text="🌐 Web Research", padding=10)
        self.web_research_panel.pack(fill=tk.X, pady=(0, 10))
        ttk.Checkbutton(self.web_research_panel, text="Enable Web Search Research", variable=self.web_research_enabled_var).pack(anchor=tk.W, pady=(0, 5))
        ttk.Label(self.web_research_panel, text="Search Method:").pack(anchor=tk.W)
        self.g_api_radio = ttk.Radiobutton(self.web_research_panel, text="🚀 Google API / DuckDuckGo (Fast)", variable=self.web_research_method_var, value="google_api")
        self.g_api_radio.pack(anchor=tk.W, padx=20)
        self.playwright_radio = ttk.Radiobutton(self.web_research_panel, text="🎭 Playwright Browser (Robust)", variable=self.web_research_method_var, value="playwright")
        self.playwright_radio.pack(anchor=tk.W, padx=20, pady=(0, 5))
        ttk.Label(self.web_research_panel, text="Max external pages to scrape from search:").pack(anchor=tk.W)
        ttk.Spinbox(self.web_research_panel, from_=1, to=20, textvariable=self.research_pages_var, width=10).pack(fill=tk.X, pady=(2, 0))
        self.yt_research_panel = ttk.LabelFrame(parent, text="📺 YouTube Research", padding=10)
        self.yt_research_panel.pack(fill=tk.X)
        self.yt_checkbutton = ttk.Checkbutton(self.yt_research_panel, text="Enable Video Transcript Analysis", variable=self.yt_research_enabled_var)
        self.yt_checkbutton.pack(anchor=tk.W, pady=(0, 5))
        ttk.Label(self.yt_research_panel, text="Videos to analyze per search query:").pack(anchor=tk.W)
        ttk.Spinbox(self.yt_research_panel, from_=1, to=5, textvariable=self.yt_videos_per_query_var, width=10).pack(fill=tk.X, pady=(2, 0))

    def create_ai_tab(self, parent):
        api_frame = ttk.LabelFrame(parent, text="🔑 API Configuration", padding=10)
        api_frame.pack(fill=tk.X, pady=(0, 10))
        ttk.Label(api_frame, text="Gemini API Key File:").pack(fill=tk.X, anchor='w')
        key_frame = ttk.Frame(api_frame)
        key_frame.pack(fill=tk.X, pady=(2, 0))
        ttk.Entry(key_frame, textvariable=self.api_key_file_var, font=("Segoe UI", 9)).pack(side=tk.LEFT, expand=True, fill=tk.X)
        ttk.Button(key_frame, text="📁", width=3, command=self.browse_api_key).pack(side=tk.RIGHT, padx=(5, 0))
        model_frame = ttk.LabelFrame(parent, text="🤖 Model Settings", padding=10)
        model_frame.pack(fill=tk.X, pady=(0, 10))
        ttk.Label(model_frame, text="Model Name:").pack(fill=tk.X, anchor='w')
        self.model_entry = ttk.Entry(model_frame, textvariable=self.model_name_var, font=("Segoe UI", 9))
        self.model_entry.pack(fill=tk.X, pady=(2, 8))
        self.temp_label = ttk.Label(model_frame, text=f"Temperature: {self.temperature_var.get():.1f}")
        self.temp_label.pack(fill=tk.X, anchor='w')
        ttk.Scale(model_frame, from_=0.0, to=1.0, orient=tk.HORIZONTAL, variable=self.temperature_var, command=self.update_temperature_label).pack(fill=tk.X, pady=(2, 8))
        ttk.Label(model_frame, text="Max Output Tokens:").pack(fill=tk.X, anchor='w')
        ttk.Spinbox(model_frame, from_=1024, to=32768, increment=1024, textvariable=self.max_tokens_var, width=10).pack(fill=tk.X, pady=(2, 0))

    def create_prompt_tab(self, parent):
        control_frame = ttk.Frame(parent); control_frame.pack(fill=tk.X, pady=(0, 10))
        ttk.Button(control_frame, text="Load Prompt", command=self.load_prompt_from_file).pack(side=tk.LEFT, expand=True, fill=tk.X, padx=(0, 5))
        ttk.Button(control_frame, text="Reset to Default", command=self.reset_prompt_to_default).pack(side=tk.RIGHT, expand=True, fill=tk.X)
        self.prompt_text = scrolledtext.ScrolledText(parent, wrap=tk.WORD, font=("Consolas", 10))
        self.prompt_text.pack(fill=tk.BOTH, expand=True, pady=(5,0))
        DarkThemeManager.configure_text_widget(self.prompt_text)

    def setup_logging(self):
        self.log_queue = queue.Queue()
        formatter = logging.Formatter('%(levelname)s: %(message)s')
        self.queue_handler = QueueHandler(self.log_queue); self.queue_handler.setFormatter(formatter)
        logging.getLogger().addHandler(self.queue_handler); logging.getLogger().setLevel(logging.INFO)
        self.root.after(100, self.poll_log_queue)

    def poll_log_queue(self):
        while True:
            try: record = self.log_queue.get(block=False)
            except queue.Empty: break
            else:
                self.log_text_widget.config(state='normal')
                tag = ''
                if "ERROR" in record or "❌" in record: tag = 'error'
                elif "WARNING" in record or "⚠️" in record: tag = 'warning'
                elif "✅" in record or "SUCCESS" in record or "🔑" in record: tag = 'success'
                elif "🔍" in record or "📺" in record or "🌐" in record or "🧠" in record or "🤖" in record: tag = 'info'
                self.log_text_widget.insert(tk.END, record + '\n', tag)
                self.log_text_widget.config(state='disabled')
                self.log_text_widget.yview(tk.END)
        self.root.after(100, self.poll_log_queue)

    def run_full_process(self):
        try:
            if not DEPENDENCIES_AVAILABLE: self.show_demo_content(); return
            self.progress_var.set("Validating inputs..."); logging.info("🚀 Starting Text-Based Study Guide Generation...")
            api_key = APIKeyManager.get_gemini_api_key(self.api_key_file_var.get())
            if not api_key:
                messagebox.showerror("API Key Error", "The Gemini API key is missing or invalid."); return
            source_data, source_name, mode = {}, "", self.input_mode_var.get()
            delay = self.config.get('scraper', {}).get('rate_limit_delay', 0.5)
            user_agent = random.choice(USER_AGENTS)
            self.progress_var.set("Collecting initial content...")
            if mode == "scrape":
                logging.info("📄 Mode: Web Scraping")
                url = self.url_var.get().strip()
                if not url: messagebox.showerror("Input Error", "URL cannot be empty."); return
                if not url.startswith(('http://', 'https://')): url = 'https://' + url; self.url_var.set(url)
                limit = self.limit_var.get()
                deep_crawl = self.deep_crawl_var.get()
                crawl_depth = self.crawl_depth_var.get()
                logging.info(f"🎯 Target: {url}"); logging.info(f"📊 Max Pages: {limit} | Deep Crawl: {deep_crawl} | Depth: {crawl_depth}")
                scraper = WebsiteScraper(url, limit, user_agent, 15, delay, gemini_api_key=api_key, deep_crawl=deep_crawl, crawl_depth=crawl_depth)
                source_data = scraper.crawl(); source_name = urlparse(url).netloc
            else: # This path now handles the TEXT-ONLY upload mode
                logging.info("📁 Mode: Local Document Processing (Text-Only)")
                file_paths = list(self.file_listbox.get(0, tk.END))
                if not file_paths: messagebox.showerror("Input Error", "Please add at least one document."); return
                logging.info(f"📚 Processing {len(file_paths)} local documents")
                loader = LocalDocumentLoader(file_paths)
                source_data = loader.load_and_extract_text(); source_name = f"{len(file_paths)}_local_documents"

            if not source_data:
                logging.warning("⚠️ No content extracted from initial source. Cannot proceed.")
                messagebox.showwarning("No Content", "No text content could be extracted from the source."); return
            logging.info(f"✅ Initial content collection complete: {len(source_data)} sources")
            if self.research_enabled_var.get():
                self.progress_var.set("Conducting AI research...")
                logging.info("\n🔬 Starting AI Research Phase"); logging.info("-" * 40)
                query_generator = EnhancedResearchQueryGenerator(api_key)
                topic_data = query_generator.extract_main_topic_and_subtopics(source_data)
                research_queries = query_generator.generate_diverse_research_queries(topic_data, self.research_queries_var.get())
                logging.info(f"🎯 Generated {len(research_queries)} research queries")
                for i, query in enumerate(research_queries, 1): logging.info(f"  {i}. {query}")
                if self.web_research_enabled_var.get():
                    self.progress_var.set("Researching web sources..."); logging.info("\n🌐 Starting Web Research")
                    exclude_domain = source_name if mode == "scrape" else ""
                    if self.web_research_method_var.get() == "playwright":
                        if PLAYWRIGHT_AVAILABLE: researcher = EnhancedPlaywrightResearcher()
                        else:
                            logging.warning("⚠️ Playwright not available, falling back to API method")
                            google_api_key, google_cx = APIKeyManager.get_google_search_credentials(self.google_api_key_file_var.get(), self.google_cx_file_var.get())
                            researcher = EnhancedGoogleSearchResearcher(google_api_key, google_cx)
                    else:
                        google_api_key, google_cx = APIKeyManager.get_google_search_credentials(self.google_api_key_file_var.get(), self.google_cx_file_var.get())
                        researcher = EnhancedGoogleSearchResearcher(google_api_key, google_cx)
                    research_urls = researcher.search_and_extract_urls(research_queries, exclude_domain)
                    if research_urls:
                        research_page_limit = self.research_pages_var.get()
                        urls_to_scrape = research_urls[:research_page_limit]
                        logging.info(f"🕷️ Scraping top {len(urls_to_scrape)} research URLs (limit: {research_page_limit})")
                        research_scraper = WebsiteScraper("http://research.local", len(urls_to_scrape), user_agent, 15, delay)
                        research_data = research_scraper.crawl(additional_urls=urls_to_scrape)
                        if research_data: source_data.update(research_data); logging.info(f"✅ Added {len(research_data)} web research sources")
                    else: logging.warning("⚠️ No research URLs found from web search")
                if self.yt_research_enabled_var.get():
                    self.progress_var.set("Analyzing YouTube videos..."); logging.info("\n📺 Starting YouTube Research")
                    if YT_DLP_AVAILABLE:
                        try:
                            yt_researcher = EnhancedYouTubeResearcher()
                            video_transcripts = yt_researcher.get_transcripts_for_queries(research_queries, self.yt_videos_per_query_var.get())
                            if video_transcripts: source_data.update(video_transcripts); logging.info(f"✅ Added {len(video_transcripts)} YouTube transcript sources")
                            else: logging.warning("⚠️ No suitable YouTube videos found")
                        except Exception as e: logging.error(f"❌ YouTube research failed: {e}")
                    else: logging.warning("⚠️ YouTube research skipped - yt-dlp not available")
            self.progress_var.set("Generating study guide..."); logging.info(f"\n📝 Starting Study Guide Generation"); logging.info("-" * 40); logging.info(f"📊 Total sources: {len(source_data)}")
            llm_config = {'model_name': self.model_name_var.get(), 'parameters': {'temperature': self.temperature_var.get(), 'max_output_tokens': self.max_tokens_var.get()}}
            generator = EnhancedNoteGenerator(api_key, llm_config, self.prompt_text.get("1.0", tk.END))
            self.final_notes_content = generator.generate_comprehensive_notes(source_data, source_name)
            self.root.after(0, lambda: self.markdown_preview.update_preview(self.final_notes_content))
            self.progress_var.set("Generation Complete!"); logging.info("\n🎉 Study Guide Generation Complete!")
            self.root.after(0, lambda: self.save_button.config(state=tk.NORMAL))
            self.root.after(0, lambda: messagebox.showinfo("Success!", f"Study guide generated successfully!\n\nSources processed: {len(source_data)}"))
        except Exception as e:
            logging.error(f"❌ Unexpected error in main process: {e}", exc_info=True)
            self.root.after(0, lambda: messagebox.showerror("Error", f"An unexpected error occurred: {e}"))
        finally:
            self.is_processing = False
            self.root.after(0, lambda: self.start_button.config(state=tk.NORMAL, text="🚀 Start Generation"))
            if not self.final_notes_content: self.root.after(0, lambda: self.progress_var.set("Ready"))
            else: self.root.after(0, lambda: self.progress_var.set("Complete - Ready to Save"))

    def run_youtube_video_analysis(self):
        logging.info("🚀 Starting Multimodal YouTube Video Analysis...")
        self.progress_var.set("Validating inputs...")
        youtube_url = self.url_var.get().strip()
        api_key_path = self.api_key_file_var.get()
        user_prompt = self.prompt_text.get("1.0", tk.END).strip()
        downloaded_video_filename = "temp_video_for_analysis.mp4"
        if not all([youtube_url, api_key_path, user_prompt]):
            messagebox.showerror("Input Error", "YouTube URL, API Key file, and a prompt are all required for this mode.")
            self.is_processing = False; self.start_button.config(state=tk.NORMAL, text="🚀 Start Generation"); return
        if not shutil.which('yt-dlp') or not shutil.which('ffmpeg'):
            logging.error("❌ 'yt-dlp' or 'ffmpeg' command not found.")
            messagebox.showerror("Missing Tools", "This feature requires 'yt-dlp' and 'ffmpeg' to be installed and in your system's PATH.")
            self.is_processing = False; self.start_button.config(state=tk.NORMAL, text="🚀 Start Generation"); return
        self.progress_var.set("Configuring API...")
        api_key = APIKeyManager.get_gemini_api_key(api_key_path)
        if not api_key:
            messagebox.showerror("API Key Error", f"Could not read a valid API key from '{api_key_path}'.");
            self.is_processing = False; self.start_button.config(state=tk.NORMAL, text="🚀 Start Generation"); return
        try:
            genai.configure(api_key=api_key)
        except Exception as e:
            logging.error(f"❌ Failed to configure Gemini API: {e}")
            messagebox.showerror("API Error", f"Failed to configure the Gemini API: {e}");
            self.is_processing = False; self.start_button.config(state=tk.NORMAL, text="🚀 Start Generation"); return
        video_file = None
        try:
            self.progress_var.set("Downloading video...")
            logging.info(f"Downloading and converting video: {youtube_url}")
            yt_dlp_command = ['yt-dlp', '-f', 'bestvideo[height<=480]+bestaudio/best[height<=480]', '--recode-video', 'mp4', '-o', downloaded_video_filename, '--force-overwrite', youtube_url]
            try:
                process = subprocess.run(yt_dlp_command, check=True, capture_output=True, text=True)
                logging.info("yt-dlp command finished successfully.")
            except subprocess.CalledProcessError as e:
                logging.error(f"yt-dlp failed. Stderr: {e.stderr}")
                messagebox.showerror("Download Error", f"yt-dlp failed to download the video.\n\nError: {e.stderr}")
                return
            if not os.path.exists(downloaded_video_filename):
                logging.error(f"Download process finished, but the file '{downloaded_video_filename}' was not created.")
                messagebox.showerror("File Error", "Video download failed. The output file was not created.")
                return
            logging.info(f"✅ Successfully created '{downloaded_video_filename}'.")
            self.progress_var.set("Uploading video to API...")
            logging.info(f"Uploading file '{downloaded_video_filename}' to the Gemini API...")
            video_file = genai.upload_file(path=downloaded_video_filename)
            logging.info(f"File uploaded. Name: {video_file.name}. Waiting for processing...")
            self.progress_var.set("Processing video...")
            while video_file.state.name == "PROCESSING":
                time.sleep(10)
                video_file = genai.get_file(name=video_file.name)
                logging.info(f"Current file state: {video_file.state.name}")
                self.progress_var.set(f"Processing... ({video_file.state.name})")
            if video_file.state.name == "FAILED":
                logging.error("Video processing failed on the server.")
                messagebox.showerror("API Error", "The Gemini API failed to process the uploaded video.")
                return
            logging.info("✅ Video processing complete.")
            self.progress_var.set("Generating analysis...")
            logging.info("Sending video and prompt to the model...")
            model_name = self.model_name_var.get() or "gemini-1.5-pro"
            logging.info(f"Using model: {model_name}")
            model = genai.GenerativeModel(model_name=model_name)
            response = model.generate_content([user_prompt, video_file])
            logging.info("✅ Model response received.")
            self.final_notes_content = response.text
            self.root.after(0, lambda: self.markdown_preview.update_preview(self.final_notes_content))
            self.progress_var.set("Generation Complete!"); logging.info("\n🎉 Multimodal Analysis Complete!")
            self.root.after(0, lambda: self.save_button.config(state=tk.NORMAL))
            self.root.after(0, lambda: messagebox.showinfo("Success!", "Multimodal video analysis completed successfully!"))
        except Exception as e:
            logging.error(f"❌ An unexpected error occurred during video analysis: {e}", exc_info=True)
            self.root.after(0, lambda: messagebox.showerror("Error", f"An unexpected error occurred: {e}"))
        finally:
            logging.info("Cleaning up resources...")
            if video_file:
                try:
                    genai.delete_file(name=video_file.name)
                    logging.info(f"Deleted remote file: {video_file.name}")
                except Exception as e: logging.warning(f"Could not delete remote file {video_file.name}: {e}")
            if os.path.exists(downloaded_video_filename):
                os.remove(downloaded_video_filename)
                logging.info(f"Deleted local file: {downloaded_video_filename}")
            self.is_processing = False
            self.root.after(0, lambda: self.start_button.config(state=tk.NORMAL, text="🚀 Start Generation"))
            if not self.final_notes_content: self.root.after(0, lambda: self.progress_var.set("Ready"))
            else: self.root.after(0, lambda: self.progress_var.set("Complete - Ready to Save"))

    ### NEW ###
    def run_multimodal_upload_analysis(self):
        logging.info("🚀 Starting Multimodal Local File Analysis...")
        self.progress_var.set("Validating inputs...")
        file_paths = list(self.file_listbox.get(0, tk.END))
        api_key_path = self.api_key_file_var.get()
        user_prompt = self.prompt_text.get("1.0", tk.END).strip()

        if not all([file_paths, api_key_path, user_prompt]):
            messagebox.showerror("Input Error", "Please add at least one file, an API Key file, and a prompt.")
            self.is_processing = False; self.start_button.config(state=tk.NORMAL, text="🚀 Start Generation"); return

        self.progress_var.set("Configuring API..."); logging.info("Configuring Gemini API...")
        api_key = APIKeyManager.get_gemini_api_key(api_key_path)
        if not api_key:
            messagebox.showerror("API Key Error", f"Could not read a valid API key from '{api_key_path}'.");
            self.is_processing = False; self.start_button.config(state=tk.NORMAL, text="🚀 Start Generation"); return
        try:
            genai.configure(api_key=api_key)
        except Exception as e:
            logging.error(f"❌ Failed to configure Gemini API: {e}")
            messagebox.showerror("API Error", f"Failed to configure the Gemini API: {e}");
            self.is_processing = False; self.start_button.config(state=tk.NORMAL, text="🚀 Start Generation"); return

        uploaded_files, text_parts, temp_image_paths = [], [], []
        temp_dir = Path(tempfile.mkdtemp())
        try:
            prompt_parts = [user_prompt]
            # --- File Processing Loop ---
            for i, file_path_str in enumerate(file_paths):
                p = Path(file_path_str)
                ext = p.suffix.lower()
                self.progress_var.set(f"Processing file {i+1}/{len(file_paths)}: {p.name}")
                logging.info(f"Processing '{p.name}'...")

                # --- Handle Images and Videos ---
                if ext in ['.png', '.jpg', '.jpeg', '.webp', '.mp4', '.mov', '.avi', '.mpeg']:
                    logging.info(f"Uploading {ext} file to API...")
                    uploaded_file = genai.upload_file(path=p)
                    uploaded_files.append(uploaded_file)
                    prompt_parts.append(uploaded_file)
                    logging.info(f"✅ File '{p.name}' uploaded and added to prompt.")

                # --- Handle PDFs (Text and Images) ---
                elif ext == '.pdf':
                    logging.info(f"Extracting text and images from PDF '{p.name}'...")
                    try:
                        doc = fitz.open(p)
                        pdf_text = ""
                        for page_num, page in enumerate(doc):
                            pdf_text += page.get_text() + "\n\n"
                            img_list = page.get_images(full=True)
                            for img_index, img in enumerate(img_list):
                                xref = img[0]
                                base_image = doc.extract_image(xref)
                                image_bytes = base_image["image"]
                                image_ext = base_image["ext"]
                                temp_img_path = temp_dir / f"{p.stem}_p{page_num}_img{img_index}.{image_ext}"
                                temp_img_path.write_bytes(image_bytes)
                                temp_image_paths.append(temp_img_path)
                        if pdf_text.strip(): text_parts.append(f"--- TEXT FROM {p.name} ---\n{pdf_text.strip()}")
                        logging.info(f"  ✅ Extracted text and {len(temp_image_paths)} images from {p.name}")
                    except Exception as e:
                        logging.error(f"  ❌ Failed to process PDF '{p.name}': {e}")

                # --- Handle Text-based files ---
                elif ext in ['.txt', '.md']: text_parts.append(f"--- CONTENT FROM {p.name} ---\n{p.read_text(encoding='utf-8', errors='ignore')}")
                elif ext == '.docx':
                    try:
                        doc_text = "\n".join(para.text for para in docx.Document(p).paragraphs)
                        text_parts.append(f"--- CONTENT FROM {p.name} ---\n{doc_text}")
                    except Exception as e: logging.error(f"  ❌ Failed to read DOCX '{p.name}': {e}")

            # --- Upload Extracted PDF Images ---
            if temp_image_paths:
                for i, img_path in enumerate(temp_image_paths):
                    self.progress_var.set(f"Uploading PDF image {i+1}/{len(temp_image_paths)}")
                    logging.info(f"Uploading extracted image: {img_path.name}")
                    uploaded_file = genai.upload_file(path=img_path)
                    uploaded_files.append(uploaded_file)
                    prompt_parts.append(uploaded_file)
                logging.info(f"✅ All {len(temp_image_paths)} extracted images uploaded.")

            # --- Combine text and add to prompt ---
            if text_parts:
                full_text_content = "\n\n".join(text_parts)
                prompt_parts.insert(1, full_text_content) # Insert text after the main prompt

            # --- Generate Content ---
            self.progress_var.set("Generating analysis..."); logging.info("Sending all content to the model...")
            model_name = self.model_name_var.get() or "gemini-1.5-pro"
            logging.info(f"Using model: {model_name}. Prompt contains {len(prompt_parts)} parts.")
            model = genai.GenerativeModel(model_name=model_name)
            response = model.generate_content(prompt_parts)
            logging.info("✅ Model response received.")

            self.final_notes_content = response.text
            self.root.after(0, lambda: self.markdown_preview.update_preview(self.final_notes_content))
            self.progress_var.set("Generation Complete!"); logging.info("\n🎉 Multimodal Analysis Complete!")
            self.root.after(0, lambda: self.save_button.config(state=tk.NORMAL))
            self.root.after(0, lambda: messagebox.showinfo("Success!", "Multimodal analysis of local files completed successfully!"))

        except Exception as e:
            logging.error(f"❌ An unexpected error occurred during multimodal analysis: {e}", exc_info=True)
            self.root.after(0, lambda: messagebox.showerror("Error", f"An unexpected error occurred: {e}"))
        finally:
            # --- Cleanup ---
            logging.info("Cleaning up resources...")
            for uploaded_file in uploaded_files:
                try:
                    logging.info(f"Deleting remote file: {uploaded_file.name}")
                    genai.delete_file(name=uploaded_file.name)
                except Exception as e: logging.warning(f"Could not delete remote file {uploaded_file.name}: {e}")
            if temp_dir.exists(): shutil.rmtree(temp_dir)
            logging.info("Cleanup complete.")
            self.is_processing = False
            self.root.after(0, lambda: self.start_button.config(state=tk.NORMAL, text="🚀 Start Generation"))
            if not self.final_notes_content: self.root.after(0, lambda: self.progress_var.set("Ready"))
            else: self.root.after(0, lambda: self.progress_var.set("Complete - Ready to Save"))


    def toggle_depth_setting(self):
        is_enabled = self.deep_crawl_var.get()
        new_state = tk.NORMAL if is_enabled else tk.DISABLED
        self.depth_label.config(state=new_state)
        self.depth_spinbox.config(state=new_state)

    def toggle_input_mode(self):
        mode = self.input_mode_var.get()
        is_scrape_or_yt = (mode in ["scrape", "youtube_video"])
        is_upload = (mode == "upload")

        # Configure URL frame
        url_frame_state = tk.NORMAL if is_scrape_or_yt else tk.DISABLED
        self._set_child_widgets_state(self.url_input_frame, url_frame_state)

        # Configure Upload frame
        upload_frame_state = tk.NORMAL if is_upload else tk.DISABLED
        self._set_child_widgets_state(self.upload_frame, upload_frame_state)

        # Specific configurations for each mode
        if mode == "scrape":
            self.url_input_frame.config(text="🕷️ Web Scraper Input")
            self.scraper_options_frame.pack(fill=tk.X, pady=(0, 0))
            self.research_checkbutton.config(state=tk.NORMAL)
            self.toggle_depth_setting()
            self.multimodal_check.config(state=tk.DISABLED)
        elif mode == "youtube_video":
            self.url_input_frame.config(text="🎬 YouTube Video (URL Analysis)")
            self.scraper_options_frame.pack_forget()
            self.research_enabled_var.set(False) # Disable external research for this mode
            self.research_checkbutton.config(state=tk.DISABLED)
            self.multimodal_check.config(state=tk.DISABLED)
        elif mode == "upload":
            self.scraper_options_frame.pack_forget()
            self.research_checkbutton.config(state=tk.NORMAL)
            self.multimodal_check.config(state=tk.NORMAL) # Enable multimodal option only for upload

        self.toggle_research_panel()

    def update_temperature_label(self, value): self.temp_label.configure(text=f"Temperature: {float(value):.1f}")
    def _set_child_widgets_state(self, parent, state):
        for widget in parent.winfo_children():
            # Skip the multimodal checkbox when disabling the parent frame so it can be controlled independently
            if widget == getattr(self, 'multimodal_check', None) and state == tk.DISABLED:
                continue
            if isinstance(widget, (ttk.Frame, ttk.LabelFrame)):
                self._set_child_widgets_state(widget, state)
            else:
                try:
                    widget.configure(state=state)
                except tk.TclError:
                    pass

    ### MODIFIED ###
    def add_files(self):
        filetypes = [
            ("All Supported Files", "*.pdf *.docx *.txt *.md *.png *.jpg *.jpeg *.webp *.mp4 *.mov *.avi"),
            ("Documents", "*.pdf *.docx *.txt *.md"),
            ("Images", "*.png *.jpg *.jpeg *.webp"),
            ("Videos", "*.mp4 *.mov *.avi *.mpeg"),
            ("All files", "*.*")
        ]
        files = filedialog.askopenfilenames(title="Select Files for Analysis", filetypes=filetypes)
        for f in files:
            if f not in self.file_listbox.get(0, tk.END):
                self.file_listbox.insert(tk.END, f)

    def clear_files(self): self.file_listbox.delete(0, tk.END)
    def browse_api_key(self):
        filepath = filedialog.askopenfilename(title="Select Gemini API Key File", filetypes=[("Key files", "*.key"), ("Text files", "*.txt")])
        if filepath: self.api_key_file_var.set(filepath)
    def load_prompt_from_file(self):
        filepath = filedialog.askopenfilename(title="Select Prompt File", filetypes=[("Markdown", "*.md"), ("Text", "*.txt")])
        if filepath:
            try:
                with open(filepath, 'r', encoding='utf-8') as f: content = f.read()
                self.prompt_text.delete("1.0", tk.END); self.prompt_text.insert(tk.END, content)
                logging.info(f"Loaded prompt from: {Path(filepath).name}")
            except Exception as e: messagebox.showerror("Error", f"Failed to load prompt: {e}")
    def toggle_research_panel(self):
        research_active = self.research_enabled_var.get()
        panel_state = tk.NORMAL if research_active else tk.DISABLED
        self._set_child_widgets_state(self.web_research_panel, panel_state)
        self._set_child_widgets_state(self.yt_research_panel, panel_state)
        if not research_active: return
        if not PLAYWRIGHT_AVAILABLE: self.playwright_radio.config(state=tk.DISABLED)
        if not YT_DLP_AVAILABLE: self.yt_checkbutton.config(state=tk.DISABLED)
    def load_initial_settings(self):
        self.config = self._load_config_file()
        prompt_template = self._load_prompt_file()
        if self.config:
            logging.info("📝 Loading settings from config.yml...")
            api_settings = self.config.get('api', {}); llm_settings = self.config.get('llm', {}); llm_params = llm_settings.get('parameters', {})
            self.api_key_file_var.set(api_settings.get('key_file', 'gemini_api.key'))
            self.model_name_var.set(llm_settings.get('model_name', 'gemini-1.5-pro'))
            self.temperature_var.set(llm_params.get('temperature', 0.5))
            self.max_tokens_var.set(llm_params.get('max_output_tokens', 8192))
            google_search_settings = api_settings.get('google_search', {})
            self.google_api_key_file_var.set(google_search_settings.get('key_file', 'google_api.key'))
            self.google_cx_file_var.set(google_search_settings.get('cx_file', 'google_cx.key'))
            logging.info("✅ Configuration loaded successfully")
        else: logging.warning("⚠️ Could not load config.yml. Using defaults.")
        if prompt_template: self.prompt_text.delete("1.0", tk.END); self.prompt_text.insert(tk.END, prompt_template)
        else: self.reset_prompt_to_default()
        self.update_temperature_label(self.temperature_var.get()); self._check_dependencies()
    def _check_dependencies(self):
        if not DEPENDENCIES_AVAILABLE: messagebox.showwarning("Missing Dependencies", "Core AI/Document libraries missing. App will run in DEMO mode.")
        if not PLAYWRIGHT_AVAILABLE:
            if hasattr(self, 'playwright_radio'): self.playwright_radio.config(state=tk.DISABLED)
            if self.web_research_method_var.get() == 'playwright': self.web_research_method_var.set('google_api')
            logging.warning("⚠️ Playwright not available. Install with: pip install playwright && playwright install")
        if not YT_DLP_AVAILABLE:
            if hasattr(self, 'yt_checkbutton'): self.yt_checkbutton.config(state=tk.DISABLED)
            self.yt_research_enabled_var.set(False)
            logging.warning("⚠️ yt-dlp not available. Install with: pip install yt-dlp")
    def reset_prompt_to_default(self):
        default_prompt = self._load_prompt_file("prompt.md") or "You are an expert educational content creator. Generate a comprehensive study guide in Markdown based on the provided content.\n\nStructure your response with:\n1. **Executive Summary**\n2. **Main Topics** (use ## and ###)\n3. **Key Points & Definitions**\n4. **Practical Examples**\n5. **Further Learning Resources**\n\nContent to analyze:\n{content}\n\nSource: {website_url}"
        self.prompt_text.delete("1.0", tk.END); self.prompt_text.insert(tk.END, default_prompt)
    def save_notes(self):
        if not self.final_notes_content: messagebox.showwarning("No Content", "No study guide to save."); return
        source_name_raw = "video_url" if self.input_mode_var.get() == "youtube_video" else (urlparse(self.url_var.get()).netloc or "website" if self.input_mode_var.get() == "scrape" else "local_files")
        source_name = re.sub(r'[^a-zA-Z0-9_-]', '', source_name_raw)
        initial_filename = f"study_guide_{source_name}_{datetime.now().strftime('%Y%m%d_%H%M')}.md"
        filepath = filedialog.asksaveasfilename(initialfile=initial_filename, defaultextension=".md", filetypes=[("Markdown", "*.md"), ("Text", "*.txt")])
        if not filepath: logging.info("💾 Save cancelled by user."); return
        try:
            with open(filepath, "w", encoding="utf-8") as f: f.write(self.final_notes_content)
            logging.info(f"💾 Study guide saved: {filepath}"); messagebox.showinfo("Success", f"Study guide saved to:\n{filepath}")
        except IOError as e: logging.error(f"❌ Failed to save file: {e}"); messagebox.showerror("Save Error", f"Could not save file: {e}")

    ### MODIFIED ###
    def start_task_thread(self):
        if self.is_processing:
            messagebox.showwarning("Processing", "A task is already running.")
            return
        self.is_processing = True
        self.start_button.config(state=tk.DISABLED, text="⏳ Processing...")
        self.save_button.config(state=tk.DISABLED)
        self.final_notes_content = ""
        self.progress_var.set("Initializing...")
        self.log_text_widget.config(state='normal')
        self.log_text_widget.delete('1.0', tk.END)
        self.log_text_widget.config(state='disabled')

        mode = self.input_mode_var.get()
        target_function = None

        # Smartly select the target function based on the mode and options
        if mode == "youtube_video":
            target_function = self.run_youtube_video_analysis
        elif mode == "upload" and self.multimodal_upload_var.get():
            target_function = self.run_multimodal_upload_analysis
        else: # Covers "scrape" mode and text-only "upload" mode
            target_function = self.run_full_process

        if target_function:
            threading.Thread(target=target_function, daemon=True).start()
        else: # Should not happen with the current logic, but good for safety
            messagebox.showerror("Error", "Could not determine the correct action to start.")
            self.is_processing = False
            self.start_button.config(state=tk.NORMAL, text="🚀 Start Generation")

    def show_demo_content(self):
        demo_content = f"# 🚀 Demo Mode\n\nThis is a demonstration because one or more critical dependencies are not installed.\n\n**Generated**: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
        self.root.after(0, lambda: self.markdown_preview.update_preview(demo_content))
        self.final_notes_content = demo_content
        self.root.after(0, lambda: self.save_button.config(state=tk.NORMAL))
        self.progress_var.set("Demo content loaded!")
    def _load_config_file(self, filepath="config.yml"):
        try:
            with open(filepath, 'r', encoding='utf-8') as f: return yaml.safe_load(f)
        except (FileNotFoundError, yaml.YAMLError): return {}
    def _load_prompt_file(self, filepath="prompt.md"):
        try:
            with open(filepath, 'r', encoding='utf-8') as f: return f.read()
        except FileNotFoundError: return ""

def main():
    """Main function to initialize and run the Tkinter application."""
    logging.basicConfig(level=logging.INFO, format='%(levelname)s - %(message)s')
    root = tk.Tk()
    app = AdvancedScraperApp(root)
    def on_closing():
        if app.is_processing and messagebox.askokcancel("Quit", "A processing task is still running. Are you sure you want to quit?"): root.destroy()
        elif not app.is_processing: root.destroy()
    root.protocol("WM_DELETE_WINDOW", on_closing)
    root.mainloop()

if __name__ == '__main__':
    main()
